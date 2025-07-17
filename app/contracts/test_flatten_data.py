import pytest
from app.contracts.service import ContractService
from pathlib import Path
from docx import Document
import os

@pytest.mark.asyncio
async def test_flatten_data_company_variants():
    service = ContractService()

    # Caso 1: company a nivel raíz con role=investor
    data1 = {
        "company": {
            "name": "GRUPO REYSA, S.R.L.",
            "rnc": "1-3225325-6",
            "rm": "3187SPM",
            "role": "investor"
        },
        "investors": [
            {
                "person": {"first_name": "Miguel Angel", "last_name": "Reyes Pichardo"},
                "person_document": {"document_number": "023-0089169-0"},
                "address": {"address_line1": "Calle Altagracia No. 39"}
            }
        ]
    }
    flat1 = service._flatten_data(data1)
    assert flat1["investor_company_name"] == "GRUPO REYSA, S.R.L."
    assert flat1["investor_rnc"] == "1-3225325-6"
    assert flat1["investor_rm"] == "3187SPM"
    assert flat1["investor_address"] == "Calle Altagracia No. 39"
    assert flat1["investor_full_name"] == "Miguel Angel Reyes Pichardo"
    assert flat1["investor_document_number"] == "023-0089169-0"

    # Caso 2: investor_company a nivel raíz
    data2 = {
        "investor_company": {
            "name": "GRUPO REYSA, S.R.L.",
            "rnc": "1-3225325-6",
            "rm": "3187SPM"
        },
        "investors": [
            {
                "person": {"first_name": "Miguel Angel", "last_name": "Reyes Pichardo"},
                "person_document": {"document_number": "023-0089169-0"},
                "address": {"address_line1": "Calle Altagracia No. 39"}
            }
        ]
    }
    flat2 = service._flatten_data(data2)
    assert flat2["investor_company_name"] == "GRUPO REYSA, S.R.L."
    assert flat2["investor_rnc"] == "1-3225325-6"
    assert flat2["investor_rm"] == "3187SPM"

    # Caso 3: company a nivel raíz con role=client
    data3 = {
        "company": {
            "name": "EMPRESA CLIENTE S.A.",
            "rnc": "1-1234567-8",
            "rm": "1234RM",
            "role": "client"
        },
        "clients": [
            {
                "person": {"first_name": "Juan", "last_name": "Pérez"},
                "person_document": {"document_number": "001-0000001-1"},
                "address": {"address_line1": "Calle Duarte #123"}
            }
        ]
    }
    flat3 = service._flatten_data(data3)
    assert flat3["client_company_name"] == "EMPRESA CLIENTE S.A."
    assert flat3["client_rnc"] == "1-1234567-8"
    assert flat3["client_rm"] == "1234RM"
    assert flat3["client_address"] == "Calle Duarte #123"
    assert flat3["client_full_name"] == "Juan Pérez"
    assert flat3["client_document"] == "001-0000001-1"

    # Caso 4: client_company a nivel raíz
    data4 = {
        "client_company": {
            "name": "EMPRESA CLIENTE S.A.",
            "rnc": "1-1234567-8",
            "rm": "1234RM"
        },
        "clients": [
            {
                "person": {"first_name": "Juan", "last_name": "Pérez"},
                "person_document": {"document_number": "001-0000001-1"},
                "address": {"address_line1": "Calle Duarte #123"}
            }
        ]
    }
    flat4 = service._flatten_data(data4)
    assert flat4["client_company_name"] == "EMPRESA CLIENTE S.A."
    assert flat4["client_rnc"] == "1-1234567-8"
    assert flat4["client_rm"] == "1234RM"

    # Caso 5: company y company en cliente/inversionista (el específico tiene prioridad)
    data5 = {
        "company": {
            "name": "NO DEBE USARSE",
            "rnc": "NO",
            "rm": "NO",
            "role": "client"
        },
        "clients": [
            {
                "company": {
                    "name": "CLIENTE PRIORITARIO S.A.",
                    "rnc": "9-9999999-9",
                    "rm": "9999RM"
                },
                "person": {"first_name": "Ana", "last_name": "López"},
                "person_document": {"document_number": "003-0000003-3"},
                "address": {"address_line1": "Calle Falsa #123"}
            }
        ]
    }
    flat5 = service._flatten_data(data5)
    assert flat5["client_company_name"] == "CLIENTE PRIORITARIO S.A."
    assert flat5["client_rnc"] == "9-9999999-9"
    assert flat5["client_rm"] == "9999RM"

@pytest.mark.asyncio
async def test_generate_contract_and_word_replacement(tmp_path):
    service = ContractService()
    data = {
        "investor_company": {
            "name": "GRUPO REYSA, S.R.L.",
            "rnc": "1-3225325-6",
            "rm": "3187SPM"
        },
        "investors": [
            {
                "person": {"first_name": "Miguel Angel", "last_name": "Reyes Pichardo"},
                "person_document": {"document_number": "023-0089169-0"},
                "address": {"address_line1": "Calle Altagracia No. 39"}
            }
        ]
    }
    # Usar una plantilla de prueba simple
    template_path = tmp_path / "test_template.docx"
    # Crear una plantilla con los placeholders
    doc = Document()
    doc.add_paragraph(
        "De una parte, la sociedad de comercio {{investor_company_name}}, organizada de acuerdo con las leyes de la República Dominicana, RNC No. {{investor_rnc}}, RM. {{investor_rm}}, con domicilio social en la {{investor_address}}, debidamente representada en este contrato por su gerente, {{investor_full_name}}, dominicano, mayor de edad, casado, portador de la cédula de identidad y electoral No.{{investor_document_number}}, con domicilio en la {{investor_address}}, sociedad que en lo que sigue del presente contrato se denominará LA PRIMERA PARTE o POR SU PROPIO NOMBRE; y,"
    )
    doc.save(template_path)
    # Forzar a usar la plantilla de prueba
    service.template_dir = tmp_path
    # Generar contrato
    result = await service.generate_contract(data)
    output_path = Path(result["path"])
    assert output_path.exists()
    # Leer el archivo generado y verificar los reemplazos
    generated_doc = Document(output_path)
    text = "\n".join([p.text for p in generated_doc.paragraphs])
    assert "GRUPO REYSA, S.R.L." in text
    assert "1-3225325-6" in text
    assert "3187SPM" in text
    assert "Calle Altagracia No. 39" in text
    assert "Miguel Angel Reyes Pichardo" in text
    assert "023-0089169-0" in text
    # No debe haber ningún {{...}} sin reemplazar
    assert "{{" not in text and "}}" not in text
    # Limpieza
    os.remove(output_path)
